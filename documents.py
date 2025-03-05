import os
import json
import pickle
import datetime
import time
from typing import List, Dict, Any, Optional, Tuple, Union
import numpy as np
import hashlib
import traceback

# Configurar a disponibilidade do LangChain
LANGCHAIN_AVAILABLE = False
try:
    import langchain
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_community.document_loaders import TextLoader, PyMuPDFLoader, Docx2txtLoader
    from langchain_core.documents import Document  # Importação corrigida
    LANGCHAIN_AVAILABLE = True
    print("LangChain disponível e será usado para processamento de documentos")
except ImportError as e:
    print(f"LangChain não está disponível: {e}")
    print("Usando implementação própria para processamento de documentos")
    LANGCHAIN_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    print("PyMuPDF não está disponível. Não será possível processar PDFs.")
    PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx não está disponível. Não será possível processar arquivos DOCX.")
    DOCX_AVAILABLE = False

# Diretório onde os documentos serão armazenados
DOCUMENTS_DIR = "documents"
# Diretório onde os chunks processados serão armazenados
CHUNKS_DIR = "document_chunks"
# Arquivo de índice para rastrear documentos
INDEX_FILE = "document_index.json"

# Criar diretórios necessários se não existirem
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(CHUNKS_DIR, exist_ok=True)

class DocumentProcessor:
    """
    Classe responsável por processar e gerenciar documentos para a base de conhecimento.
    """
    
    def __init__(self):
        """Inicializa o processador de documentos e carrega o índice."""
        self.documents_index = self._load_index()
        
    def _load_index(self) -> Dict[str, Any]:
        """Carrega o índice de documentos do arquivo JSON."""
        if os.path.exists(INDEX_FILE):
            try:
                with open(INDEX_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"Erro ao carregar índice: {e}")
                return {}
        return {}
    
    def _save_index(self) -> None:
        """Salva o índice de documentos no arquivo JSON."""
        try:
            with open(INDEX_FILE, 'w', encoding='utf-8') as f:
                json.dump(self.documents_index, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar índice: {e}")
            
    def get_document_list(self) -> List[Dict[str, Any]]:
        """Retorna a lista de documentos disponíveis no índice."""
        return list(self.documents_index.values())
    
    def add_document(self, file_path: str, title: Optional[str] = None, description: Optional[str] = None) -> str:
        """
        Adiciona um novo documento à base de conhecimento.
        
        Args:
            file_path: Caminho do arquivo a ser adicionado
            title: Título do documento (opcional)
            description: Descrição do documento (opcional)
            
        Returns:
            ID do documento adicionado
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        # Obter extensão e verificar se é suportada
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf' and not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF (fitz) é necessário para processar PDFs. Instale com: pip install pymupdf")
        
        if file_ext == '.docx' and not DOCX_AVAILABLE:
            raise ImportError("python-docx é necessário para processar arquivos DOCX. Instale com: pip install python-docx")
        
        # Criar ID único com base no conteúdo e nome do arquivo
        with open(file_path, 'rb') as f:
            file_content = f.read()
            doc_id = hashlib.md5(file_content + os.path.basename(file_path).encode()).hexdigest()
        
        # Nome do arquivo de destino
        dest_filename = f"{doc_id}{file_ext}"
        dest_path = os.path.join(DOCUMENTS_DIR, dest_filename)
        
        # Copiar arquivo para o diretório de documentos
        with open(file_path, 'rb') as src, open(dest_path, 'wb') as dst:
            dst.write(src.read())
            
        # Extrair texto do documento
        text = self._extract_text_from_file(dest_path)
        
        # Processar e segmentar o texto
        chunks = self._split_text(text, doc_id)
        
        # Salvar os chunks
        self._save_chunks(chunks, doc_id)
        
        # Adicionar documento ao índice
        if not title:
            title = os.path.basename(file_path)
            
        document_info = {
            "id": doc_id,
            "title": title,
            "description": description or "",
            "filename": dest_filename,
            "original_filename": os.path.basename(file_path),
            "file_type": file_ext[1:],  # remover o ponto
            "added_date": datetime.datetime.now().isoformat(),
            "num_chunks": len(chunks)
        }
        
        self.documents_index[doc_id] = document_info
        self._save_index()
        
        return doc_id
    
    def remove_document(self, doc_id: str) -> bool:
        """
        Remove um documento da base de conhecimento.
        
        Args:
            doc_id: ID do documento a ser removido
            
        Returns:
            True se o documento foi removido com sucesso, False caso contrário
        """
        try:
            # Verificar se o documento existe no índice
            if doc_id not in self.documents_index:
                print(f"Documento com ID {doc_id} não encontrado no índice")
                return False
            
            # Obter informações do documento
            doc_info = self.documents_index[doc_id]
            print(f"Removendo documento: {doc_info['title']} (ID: {doc_id})")
            
            # Caminhos dos arquivos a serem removidos
            doc_file = os.path.join(DOCUMENTS_DIR, doc_info['filename'])
            chunks_file = os.path.join(CHUNKS_DIR, f"{doc_id}.json")
            
            # Remover arquivo do documento (se existir)
            if os.path.exists(doc_file):
                try:
                    os.remove(doc_file)
                    print(f"Arquivo do documento removido: {doc_file}")
                except Exception as e:
                    print(f"Erro ao remover arquivo do documento: {str(e)}")
            else:
                print(f"Arquivo do documento não encontrado: {doc_file}")
            
            # Remover arquivo de chunks (se existir)
            if os.path.exists(chunks_file):
                try:
                    os.remove(chunks_file)
                    print(f"Arquivo de chunks removido: {chunks_file}")
                except Exception as e:
                    print(f"Erro ao remover arquivo de chunks: {str(e)}")
            else:
                print(f"Arquivo de chunks não encontrado: {chunks_file}")
            
            # Remover documento do índice
            del self.documents_index[doc_id]
            self._save_index()
            print(f"Documento removido do índice com sucesso")
            
            return True
            
        except Exception as e:
            import traceback
            print(f"Erro ao remover documento: {str(e)}")
            traceback.print_exc()
            return False
    
    def reprocess_document(self, doc_id: str) -> Tuple[bool, str]:
        """
        Reprocessa um documento existente para atualizar seus chunks.
        
        Args:
            doc_id: ID do documento a ser reprocessado
            
        Returns:
            Tupla (sucesso, mensagem)
        """
        try:
            # Verificar se o documento existe no índice
            if doc_id not in self.documents_index:
                return False, f"Documento com ID {doc_id} não encontrado"
            
            # Obter informações do documento
            doc_info = self.documents_index[doc_id]
            print(f"Reprocessando documento: {doc_info['title']} (ID: {doc_id})")
            
            # Caminho do arquivo do documento
            doc_file = os.path.join(DOCUMENTS_DIR, doc_info['filename'])
            
            # Verificar se o arquivo existe
            if not os.path.exists(doc_file):
                return False, f"Arquivo do documento não encontrado: {doc_file}"
            
            # Extrair texto do documento
            try:
                text = self._extract_text_from_file(doc_file)
            except Exception as e:
                return False, f"Erro ao extrair texto: {str(e)}"
            
            # Processar o texto em chunks
            chunks = self._split_text(text, doc_id)
            
            # Salvar os novos chunks
            self._save_chunks(chunks, doc_id)
            
            return True, f"Documento '{doc_info['title']}' reprocessado com sucesso!"
            
        except Exception as e:
            print(f"Erro ao reprocessar documento: {str(e)}")
            return False, f"Erro ao reprocessar documento: {str(e)}"
    
    def reprocess_all_documents(self) -> Tuple[bool, str]:
        """
        Reprocessa todos os documentos, gerando embeddings para todos os chunks.
        
        Returns:
            Tupla com (sucesso, mensagem)
        """
        try:
            # Listar todos os documentos no índice
            docs = list(self.documents_index.keys())
            
            if not docs:
                return False, "Nenhum documento encontrado no índice"
            
            print(f"Reprocessando {len(docs)} documentos...")
            
            # Reprocessar cada documento
            for doc_id in docs:
                doc_info = self.documents_index[doc_id]
                doc_path = doc_info.get("path", "")
                
                if not doc_path or not os.path.exists(doc_path):
                    print(f"Caminho não encontrado para documento {doc_id}: {doc_path}")
                    continue
                    
                print(f"Reprocessando documento: {doc_info.get('name', doc_id)}")
                
                # Reprocessar o documento
                self.reprocess_document(doc_id)
            
            return True, f"Reprocessados {len(docs)} documentos com sucesso"
            
        except Exception as e:
            error_msg = f"Erro ao reprocessar documentos: {str(e)}"
            print(error_msg)
            traceback.print_exc()
            return False, error_msg
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """
        Extrai o texto de um arquivo de acordo com seu tipo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Texto extraído do arquivo
        """
        # Verificar se o arquivo existe
        if not os.path.exists(file_path):
            return f"ERRO: Arquivo não encontrado: {file_path}"
        
        # Verificar se o LangChain está disponível
        try:
            from langchain_community.vectorstores import FAISS
            from langchain_openai import OpenAIEmbeddings
            from langchain_core.documents import Document
            LANGCHAIN_AVAILABLE = True
            print("LangChain está disponível e importado com sucesso!")
        except ImportError as e:
            print(f"LangChain não está disponível: {e}")
            LANGCHAIN_AVAILABLE = False
        except Exception as e:
            print(f"Erro ao importar LangChain: {e}")
            LANGCHAIN_AVAILABLE = False
        
        # Usar LangChain se disponível
        if LANGCHAIN_AVAILABLE:
            try:
                print(f"Usando LangChain para extrair texto de {file_path}")
                file_ext = os.path.splitext(file_path)[1].lower()
                
                if file_ext == '.pdf':
                    loader = PyMuPDFLoader(file_path)
                    documents = loader.load()
                    return "\n\n".join([doc.page_content for doc in documents])
                
                elif file_ext == '.docx':
                    loader = Docx2txtLoader(file_path)
                    documents = loader.load()
                    return "\n\n".join([doc.page_content for doc in documents])
                
                elif file_ext == '.txt' or file_ext == '.md':
                    loader = TextLoader(file_path, encoding='utf-8')
                    documents = loader.load()
                    return "\n\n".join([doc.page_content for doc in documents])
                
                else:
                    print(f"LangChain: Tipo de arquivo não suportado: {file_ext}")
                    # Fallback para método padrão
            
            except Exception as e:
                print(f"Erro ao usar LangChain para extrair texto: {str(e)}")
                print("Voltando ao método padrão de extração")
        
        # Método padrão se LangChain não estiver disponível ou falhar
        print(f"Usando método padrão para extrair texto de {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf' and PYMUPDF_AVAILABLE:
            # Para arquivos PDF, usa PyMuPDF
            text = ""
            try:
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text() + "\n\n"
                return text
            except Exception as e:
                print(f"Erro ao processar PDF: {str(e)}")
                return f"ERRO AO PROCESSAR PDF: {str(e)}"
        
        elif file_ext == '.txt' or file_ext == '.md':
            # Para arquivos de texto, lê diretamente
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Se falhar com UTF-8, tenta com latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                print(f"Erro ao ler arquivo de texto: {str(e)}")
                return f"ERRO AO LER ARQUIVO: {str(e)}"
        
        elif file_ext == '.docx' and DOCX_AVAILABLE:
            # Para arquivos DOCX, usa python-docx
            text = ""
            try:
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except Exception as e:
                print(f"Erro ao processar DOCX: {str(e)}")
                return f"ERRO AO PROCESSAR DOCX: {str(e)}"
            
        else:
            # Para outros tipos de arquivo, retorna uma mensagem
            return f"Tipo de arquivo não suportado: {file_ext}"
    
    def _split_text(self, text: str, doc_id: str) -> List[Dict[str, Any]]:
        """
        Divide o texto em chunks menores para processamento.
        
        Args:
            text: Texto completo do documento
            doc_id: ID do documento
            
        Returns:
            Lista de chunks processados
        """
        if LANGCHAIN_AVAILABLE:
            # Usar LangChain para dividir o texto
            try:
                print(f"Usando LangChain para dividir o texto do documento {doc_id}")
                # Criar um text splitter com configurações otimizadas
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=100,
                    length_function=len,
                    separators=["\n\n", "\n", " ", ""]
                )
                
                # Dividir o texto em chunks
                langchain_chunks = text_splitter.split_text(text)
                
                # Converter para o nosso formato de chunks
                chunks = []
                for i, chunk_text in enumerate(langchain_chunks):
                    chunk = {
                        "id": f"{doc_id}_chunk_{i+1}",
                        "doc_id": doc_id,
                        "text": chunk_text.strip(),
                        "embedding": None
                    }
                    chunks.append(chunk)
                
                print(f"LangChain dividiu o texto em {len(chunks)} chunks")
                return chunks
            except Exception as e:
                print(f"Erro ao usar LangChain para dividir texto: {str(e)}")
                print("Voltando ao método de divisão padrão")
                # Em caso de erro, voltar ao método padrão
        
        # Método padrão se LangChain não estiver disponível ou falhar
        print(f"Usando método padrão para dividir o texto do documento {doc_id}")
        chunks = []
        chunk_id = 1
        
        # Método alternativo simples de divisão de texto
        # Dividir por parágrafos e depois combinar para chunks de tamanho aproximado
        paragraphs = text.split("\n\n")
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < 1000:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunk = {
                        "id": f"{doc_id}_chunk_{chunk_id}",
                        "doc_id": doc_id,
                        "text": current_chunk.strip(),
                        "embedding": None
                    }
                    chunks.append(chunk)
                    chunk_id += 1
                current_chunk = para + "\n\n"
        
        # Adicionar o último chunk se houver conteúdo
        if current_chunk:
            chunk = {
                "id": f"{doc_id}_chunk_{chunk_id}",
                "doc_id": doc_id,
                "text": current_chunk.strip(),
                "embedding": None
            }
            chunks.append(chunk)
        
        print(f"Método padrão dividiu o texto em {len(chunks)} chunks")
        return chunks
    
    def _save_chunks(self, chunks: List[Dict[str, Any]], doc_id: str) -> None:
        """
        Salva os chunks de um documento em um arquivo JSON.
        
        Args:
            chunks: Lista de chunks a serem salvos
            doc_id: ID do documento
        """
        chunks_path = os.path.join(CHUNKS_DIR, f"{doc_id}.json")
        try:
            with open(chunks_path, 'w', encoding='utf-8') as f:
                json.dump(chunks, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar chunks: {e}")
    
    def get_document_chunks(self, doc_id: str) -> List[Dict[str, Any]]:
        """
        Recupera os chunks de um documento.
        
        Args:
            doc_id: ID do documento
            
        Returns:
            Lista de chunks do documento
        """
        chunks_path = os.path.join(CHUNKS_DIR, f"{doc_id}.json")
        if not os.path.exists(chunks_path):
            return []
            
        try:
            with open(chunks_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Erro ao carregar chunks: {e}")
            return []
            
    def search_documents(self, query: str) -> List[Dict[str, Any]]:
        """
        Realiza uma busca nos documentos usando embeddings para comparação semântica.
        
        Args:
            query: Termo de busca
            
        Returns:
            Lista de chunks que correspondem à busca, ordenados por relevância
        """
        print(f"Realizando busca com embeddings para: '{query}'")
        query = query.lower()
        results = []
        
        try:
            # Gerar embedding para a consulta
            query_embedding = generate_embedding(query)
            
            # Buscar em todos os documentos
            for doc_id in self.documents_index:
                chunks = self.get_document_chunks(doc_id)
                
                for chunk in chunks:
                    # Verificar se o chunk tem embedding
                    chunk_embedding = chunk.get("embedding", [])
                    
                    if chunk_embedding:
                        # Calcular similaridade
                        similarity = cosine_similarity(query_embedding, chunk_embedding)
                        
                        # Adicionar à lista de resultados se tiver similaridade mínima
                        if similarity > 0.05:  # Reduzido de 0.1 para 0.05 para ser mais permissivo
                            # Adicionar informações do documento ao chunk
                            doc_info = self.documents_index[doc_id].copy()
                            chunk_with_doc = chunk.copy()
                            chunk_with_doc["document"] = doc_info
                            chunk_with_doc["similarity"] = similarity
                            results.append(chunk_with_doc)
                    else:
                        # Fallback para busca textual se o chunk não tiver embedding
                        chunk_text = chunk["text"].lower()
                        
                        # Busca textual mais avançada
                        # 1. Verificar correspondência exata
                        if query in chunk_text:
                            doc_info = self.documents_index[doc_id].copy()
                            chunk_with_doc = chunk.copy()
                            chunk_with_doc["document"] = doc_info
                            chunk_with_doc["similarity"] = 0.5  # Valor médio para correspondências de texto exato
                            results.append(chunk_with_doc)
                        else:
                            # 2. Verificar correspondência parcial de palavras
                            query_words = query.split()
                            matched_words = sum(1 for word in query_words if len(word) > 3 and word in chunk_text)
                            
                            # Aceitar correspondências parciais
                            if matched_words > 0:
                                similarity_score = 0.3 * (matched_words / len(query_words))
                                doc_info = self.documents_index[doc_id].copy()
                                chunk_with_doc = chunk.copy()
                                chunk_with_doc["document"] = doc_info
                                chunk_with_doc["similarity"] = similarity_score
                                results.append(chunk_with_doc)
                                print(f"Fallback - correspondência parcial: {matched_words}/{len(query_words)} palavras")
            
            # Ordenar por similaridade (maior para menor)
            results.sort(key=lambda x: x.get("similarity", 0), reverse=True)
            
            print(f"Encontrados {len(results)} resultados com embeddings")
            
        except Exception as e:
            print(f"Erro na busca com embeddings: {str(e)}")
            print("Fallback para busca textual simples")
            
            # Fallback para busca textual simples em caso de erro
            for doc_id in self.documents_index:
                chunks = self.get_document_chunks(doc_id)
                
                for chunk in chunks:
                    chunk_text = chunk["text"].lower()
                    
                    # 1. Verificar correspondência exata
                    if query in chunk_text:
                        doc_info = self.documents_index[doc_id].copy()
                        chunk_with_doc = chunk.copy()
                        chunk_with_doc["document"] = doc_info
                        chunk_with_doc["similarity"] = 0.5
                        results.append(chunk_with_doc)
                    else:
                        # 2. Verificar correspondência parcial de palavras
                        query_words = query.split()
                        matched_words = sum(1 for word in query_words if len(word) > 3 and word in chunk_text)
                        
                        # Aceitar correspondências parciais
                        if matched_words > 0:
                            similarity_score = 0.3 * (matched_words / len(query_words))
                            doc_info = self.documents_index[doc_id].copy()
                            chunk_with_doc = chunk.copy()
                            chunk_with_doc["document"] = doc_info
                            chunk_with_doc["similarity"] = similarity_score
                            results.append(chunk_with_doc)
                            print(f"Fallback - correspondência parcial: {matched_words}/{len(query_words)} palavras")
        
        # Se não encontramos nenhuma correspondência, pegar alguns chunks de cada documento
        if not results:
            print("Nenhum resultado encontrado. Adicionando chunks de fallback.")
            for doc_id in self.documents_index:
                chunks = self.get_document_chunks(doc_id)
                
                # Pegar os 2 primeiros chunks de cada documento
                fallback_chunks = chunks[:2]
                for chunk in fallback_chunks:
                    doc_info = self.documents_index[doc_id].copy()
                    chunk_with_doc = chunk.copy()
                    chunk_with_doc["document"] = doc_info
                    results.append(chunk_with_doc)
                    
        return results

# Função para obter uma instância única do processador de documentos
def get_document_processor() -> DocumentProcessor:
    """Retorna uma instância do processador de documentos."""
    return DocumentProcessor()

# Função para facilitar a adição de documentos via interface
def add_document_from_upload(uploaded_file, title=None, description=None) -> Tuple[bool, str]:
    """
    Adiciona um documento enviado através de um upload.
    
    Args:
        uploaded_file: Arquivo enviado (normalmente de st.file_uploader)
        title: Título do documento
        description: Descrição do documento
        
    Returns:
        Tupla (sucesso, mensagem)
    """
    try:
        # Verificar se o arquivo foi enviado
        if uploaded_file is None:
            return False, "Nenhum arquivo enviado"
            
        # Verificar o tipo de arquivo
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        
        # Verificar dependências
        if file_ext == '.pdf':
            try:
                import fitz
                print("PyMuPDF disponível para processamento")
            except ImportError:
                return False, "PyMuPDF (fitz) é necessário para processar PDFs. Instale com: pip install pymupdf"
                
        if file_ext == '.docx':
            try:
                import docx
                print("python-docx disponível para processamento")
            except ImportError:
                return False, "python-docx é necessário para processar arquivos DOCX. Instale com: pip install python-docx"
        
        # Criar diretórios se não existirem
        os.makedirs(DOCUMENTS_DIR, exist_ok=True)
        os.makedirs(CHUNKS_DIR, exist_ok=True)
        
        # Usar um nome de arquivo sem espaços para evitar problemas
        safe_filename = uploaded_file.name.replace(" ", "_")
        temp_file_path = os.path.join(DOCUMENTS_DIR, f"temp_{safe_filename}")
        
        # Salvar o arquivo enviado
        print(f"Salvando arquivo temporário: {temp_file_path}")
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
            
        # Verificar se o arquivo foi salvo corretamente
        if not os.path.exists(temp_file_path):
            return False, "Falha ao salvar o arquivo temporário"
            
        file_size = os.path.getsize(temp_file_path)
        print(f"Arquivo salvo com {file_size} bytes")
        
        if file_size == 0:
            os.remove(temp_file_path)
            return False, "O arquivo foi salvo, mas está vazio"
        
        # Processar o documento
        try:
            processor = get_document_processor()
            print(f"Processando documento: {temp_file_path}")
            
            # Adicionar o documento
            doc_id = processor.add_document(temp_file_path, title, description)
            print(f"Documento processado com ID: {doc_id}")
            
            # Remover o arquivo temporário após o processamento
            try:
                os.remove(temp_file_path)
                print(f"Arquivo temporário removido: {temp_file_path}")
            except Exception as e:
                print(f"Aviso: Não foi possível remover o arquivo temporário: {str(e)}")
            
            # Verificar se os chunks foram criados
            chunks_path = os.path.join(CHUNKS_DIR, f"{doc_id}.json")
            if os.path.exists(chunks_path):
                return True, f"Documento '{title}' adicionado com sucesso!"
            else:
                return False, "O documento foi processado, mas os chunks não foram gerados."
                
        except Exception as e:
            import traceback
            error_msg = f"Erro ao processar documento: {str(e)}"
            print(error_msg)
            traceback.print_exc()
            
            # Tentar remover o arquivo temporário em caso de erro
            try:
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
            except:
                pass
                
            return False, error_msg
            
    except Exception as e:
        import traceback
        error_msg = f"Erro ao adicionar documento: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        return False, error_msg

# Função para reconstruir o índice a partir dos arquivos existentes
def rebuild_document_index() -> bool:
    """
    Reconstrói o índice de documentos com base nos arquivos existentes.
    Isso é útil quando o arquivo de índice está corrompido ou quando há 
    inconsistências entre o índice e os arquivos.
    
    Returns:
        True se o índice foi reconstruído com sucesso, False caso contrário
    """
    try:
        # Novos índices vazios
        new_index = {}
        
        # Verificar documentos existentes
        if not os.path.exists(DOCUMENTS_DIR):
            print(f"Diretório de documentos não existe: {DOCUMENTS_DIR}")
            return False
            
        # Verificar chunks existentes
        if not os.path.exists(CHUNKS_DIR):
            print(f"Diretório de chunks não existe: {CHUNKS_DIR}")
            return False
            
        # Listar arquivos de chunks (eles contêm os IDs de documento)
        chunk_files = [f for f in os.listdir(CHUNKS_DIR) if f.endswith('.json')]
        
        # Para cada arquivo de chunks
        for chunk_file in chunk_files:
            # Extrair o ID do documento do nome do arquivo (removendo a extensão .json)
            doc_id = os.path.splitext(chunk_file)[0]
            
            # Procurar o arquivo do documento correspondente
            document_files = os.listdir(DOCUMENTS_DIR)
            matching_docs = [f for f in document_files if f.startswith(doc_id)]
            
            if matching_docs:
                doc_filename = matching_docs[0]
                file_ext = os.path.splitext(doc_filename)[1].lower()
                
                # Verificar número de chunks
                chunks_path = os.path.join(CHUNKS_DIR, chunk_file)
                try:
                    with open(chunks_path, 'r', encoding='utf-8') as f:
                        chunks = json.load(f)
                        num_chunks = len(chunks)
                except:
                    num_chunks = 0
                
                # Criar entrada de índice
                new_index[doc_id] = {
                    "id": doc_id,
                    "title": f"Documento {doc_id[:6]}",
                    "description": "Documento restaurado automaticamente",
                    "filename": doc_filename,
                    "original_filename": doc_filename,
                    "file_type": file_ext[1:],  # remover o ponto
                    "added_date": datetime.datetime.now().isoformat(),
                    "num_chunks": num_chunks
                }
        
        # Salvar o novo índice
        with open(INDEX_FILE, 'w', encoding='utf-8') as f:
            json.dump(new_index, f, ensure_ascii=False, indent=2)
            
        print(f"Índice reconstruído com {len(new_index)} documentos")
        return True
        
    except Exception as e:
        import traceback
        print(f"Erro ao reconstruir índice: {str(e)}")
        traceback.print_exc()
        return False

# Função para buscar conhecimento relevante
def get_relevant_knowledge(question: str, max_chunks: int = 10, max_chars: int = 5000, similarity_threshold: float = 0.2) -> str:
    """
    Busca conhecimento relevante nos documentos para responder à pergunta.
    
    Args:
        question: Pergunta do usuário
        max_chunks: Número máximo de chunks a serem retornados
        max_chars: Número máximo de caracteres total
        similarity_threshold: Limiar mínimo de similaridade (valores menores são mais permissivos)
        
    Returns:
        String contendo o conhecimento relevante
    """
    print(f"\n=== Buscando conhecimento relevante para: '{question}' ===")
    
    # Normalizar a pergunta (remover acentuação, converter para minúsculas)
    import unicodedata
    normalized_question = unicodedata.normalize('NFKD', question.lower())
    normalized_question = ''.join([c for c in normalized_question if not unicodedata.combining(c)])
    print(f"Pergunta normalizada: '{normalized_question}'")
    
    # Verificar se podemos usar o LangChain
    if LANGCHAIN_AVAILABLE:
        try:
            print("Tentando usar LangChain para busca de conhecimento relevante")
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                print("AVISO: OPENAI_API_KEY não está configurada para LangChain")
                # Fallback para método padrão
            else:
                # Verificar documentos disponíveis
                processor = get_document_processor()
                document_list = processor.get_document_list()
                
                if not document_list:
                    return "Não foram encontrados documentos na base de conhecimento."
                
                # Tentar construir um índice de vetores com FAISS
                from langchain_community.vectorstores import FAISS  
                from langchain_core.documents import Document  
                
                # Coletar documentos para indexar
                documents = []
                for doc_info in document_list:
                    doc_id = doc_info["id"]
                    chunks_file = os.path.join(processor.chunks_dir, f"{doc_id}.json")
                    
                    if os.path.exists(chunks_file):
                        with open(chunks_file, 'r', encoding='utf-8') as f:
                            doc_chunks = json.load(f)
                            
                        for chunk in doc_chunks:
                            # Criar documento LangChain
                            langchain_doc = Document(
                                page_content=chunk["text"],
                                metadata={
                                    "doc_id": doc_id,
                                    "chunk_id": chunk["id"],
                                    "title": doc_info["title"]
                                }
                            )
                            documents.append(langchain_doc)
                
                if not documents:
                    return "Nenhum chunk de documento encontrado."
                
                print(f"Criando índice FAISS com {len(documents)} documentos")
                
                # Criar embeddings
                embeddings = OpenAIEmbeddings(openai_api_key=api_key)
                
                # Criar vectorstore
                try:
                    db = FAISS.from_documents(documents, embeddings)
                    
                    # Realizar busca com mais resultados e menor limiar de similaridade
                    k_search = max_chunks * 2  # Buscar mais resultados para ter mais opções
                    similar_docs = db.similarity_search_with_score(normalized_question, k=k_search)
                    
                    # Construir resposta
                    knowledge = ""
                    for i, (doc, score) in enumerate(similar_docs):
                        title = doc.metadata.get("title", "Desconhecido")
                        chunk_text = doc.page_content
                        
                        # Verificar limite de caracteres
                        if len(knowledge) + len(chunk_text) > max_chars:
                            # Se o chunk for muito grande, tentar cortar para encaixar no limite
                            remaining_space = max_chars - len(knowledge)
                            if remaining_space > 100:  # Se tiver pelo menos 100 caracteres de espaço
                                chunk_text = chunk_text[:remaining_space] + "..."
                            else:
                                break
                        
                        knowledge += f"\n--- Início do trecho {i+1} (de {title}, similaridade: {1-score:.4f}) ---\n"
                        knowledge += chunk_text
                        knowledge += f"\n--- Fim do trecho {i+1} ---\n"
                        
                        # Logging
                        preview = chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text
                        print(f"LangChain - Chunk {i+1}: {preview}")
                        print(f"  - Documento: {title}")
                        print(f"  - Similaridade: {1-score:.4f}")
                        print(f"  - Tamanho: {len(chunk_text)} caracteres")
                    
                    if knowledge:
                        print(f"LangChain encontrou {len(similar_docs)} chunks relevantes")
                        return knowledge
                    else:
                        return "LangChain não encontrou informações relevantes."
                except ImportError as e:
                    print(f"Erro ao usar FAISS: {e}")
                    print("A biblioteca FAISS não está disponível, usando método padrão")
                    # Continuar com o método padrão
        except Exception as e:
            print(f"Erro ao usar LangChain para busca: {str(e)}")
            traceback.print_exc()
            print("Voltando ao método padrão de busca")
    
    # Método padrão se LangChain não estiver disponível ou falhar
    # Obter o processador de documentos
    processor = get_document_processor()
    
    # Buscar documentos relevantes
    relevant_chunks = processor.search_documents(normalized_question)
    
    # Limitar o número de chunks para evitar exceder o contexto do LLM
    if len(relevant_chunks) > max_chunks:
        print(f"Limitando de {len(relevant_chunks)} para {max_chunks} chunks")
        relevant_chunks = relevant_chunks[:max_chunks]
    
    print(f"Encontrados {len(relevant_chunks)} chunks relevantes")
    
    # Construir o contexto
    knowledge = ""
    total_chars = 0
    
    # Logging para ver quais chunks estão sendo enviados
    print(f"\nChunks selecionados para a pergunta: '{question}'")
    
    for i, chunk in enumerate(relevant_chunks):
        chunk_text = chunk["text"]
        doc_info = chunk.get("document", {})
        doc_name = doc_info.get("name", "Desconhecido")
        doc_title = doc_info.get("title", doc_name)
        similarity = chunk.get("similarity", 0)
        
        # Verificar o limite de caracteres
        if total_chars + len(chunk_text) > max_chars:
            # Se o chunk for muito grande, tentar cortar para encaixar no limite
            remaining_space = max_chars - total_chars
            if remaining_space > 100:  # Se tiver pelo menos 100 caracteres de espaço
                chunk_text = chunk_text[:remaining_space] + "..."
            else:
                break
                
        # Adicionar o chunk ao conhecimento
        knowledge += f"\n--- Início do trecho {i+1} (de {doc_title}) ---\n"
        knowledge += chunk_text
        knowledge += f"\n--- Fim do trecho {i+1} ---\n"
        
        # Atualizar contagem de caracteres
        total_chars += len(chunk_text)
        
        # Logging
        preview = chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text
        print(f"Chunk {i+1}: {preview}")
        print(f"  - Documento: {doc_title}")
        print(f"  - Similaridade: {similarity:.4f}")
        print(f"  - Tamanho: {len(chunk_text)} caracteres")
        
    print(f"\nTotal de caracteres enviados ao LLM: {total_chars}")
    
    if not knowledge:
        knowledge = "Não foram encontradas informações relevantes nos documentos."
    
    return knowledge

# Função para gerar embeddings usando a API da OpenAI
def generate_embedding(text: str) -> List[float]:
    """
    Gera um embedding para um texto usando a API da OpenAI.
    
    Args:
        text: Texto para gerar embedding
        
    Returns:
        Lista de floats representando o embedding ou lista vazia em caso de erro
    """
    if not text.strip():
        print("Texto vazio, retornando embedding vazio")
        return []
    
    # Verificar se podemos usar o LangChain para embeddings
    if LANGCHAIN_AVAILABLE:
        try:
            print("Usando LangChain para gerar embedding")
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                print("AVISO: OPENAI_API_KEY não está configurada")
                return []
                
            embeddings = OpenAIEmbeddings(openai_api_key=api_key)
            embedding_vector = embeddings.embed_query(text)
            return embedding_vector
        except Exception as e:
            print(f"Erro ao usar LangChain para embedding: {str(e)}")
            print("Voltando ao método padrão para embeddings")
    
    # Método padrão se LangChain não estiver disponível ou falhar
    max_retries = 3
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            # Importa openai aqui para não quebrar na importação inicial
            import openai
            from openai import OpenAI
            
            # Tenta usar a nova sintaxe primeiro
            try:
                # Verificar se a chave API está configurada
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                    print("AVISO: OPENAI_API_KEY não está configurada")
                    return []
                
                client = OpenAI(api_key=api_key)
                response = client.embeddings.create(
                    input=text,
                    model="text-embedding-ada-002"
                )
                return response.data[0].embedding
            except (AttributeError, ImportError):
                # Fallback para a sintaxe antiga
                response = openai.Embedding.create(
                    input=text,
                    model="text-embedding-ada-002"
                )
                return response["data"][0]["embedding"]
        except Exception as e:
            retry_count += 1
            print(f"Erro ao gerar embedding (tentativa {retry_count}/{max_retries}): {str(e)}")
            import time
            time.sleep(1)  # Esperar um pouco antes de tentar novamente
    
    print("Falha em todas as tentativas de gerar embedding. Retornando lista vazia.")
    return []

def cosine_similarity(a: List[float], b: List[float]) -> float:
    """
    Calcula a similaridade de cosseno entre dois vetores.
    
    Args:
        a: Primeiro vetor
        b: Segundo vetor
        
    Returns:
        Similaridade de cosseno (0-1, onde 1 é mais similar)
    """
    if not a or not b:
        return 0.0
        
    a = np.array(a)
    b = np.array(b)
    
    # Normalizar os vetores
    a_norm = np.linalg.norm(a)
    b_norm = np.linalg.norm(b)
    
    # Evitar divisão por zero
    if a_norm == 0 or b_norm == 0:
        return 0.0
        
    return np.dot(a, b) / (a_norm * b_norm)

# Função para reprocessar todos os documentos
def reprocess_all_documents() -> Tuple[bool, str]:
    """
    Reprocessa todos os documentos, gerando embeddings para todos os chunks.
    
    Returns:
        Tupla com (sucesso, mensagem)
    """
    try:
        processor = get_document_processor()
        return processor.reprocess_all_documents()
    except Exception as e:
        error_msg = f"Erro ao reprocessar documentos: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        return False, error_msg

# Exemplo de uso
if __name__ == "__main__":
    processor = get_document_processor()
    print("Documentos carregados:", len(processor.get_document_list()))
    
    # Exemplo de como adicionar um documento (descomente para testar)
    # doc_id = processor.add_document("caminho/para/documento.pdf", "Título do Documento", "Descrição opcional")
    # print(f"Documento adicionado: {doc_id}")
    
    # Exemplo de busca (descomente para testar)
    # results = processor.search_documents("termo de busca")
    # print(f"Resultados encontrados: {len(results)}")